"""Step 7: Report Generation."""

import streamlit as st
import pandas as pd
from io import BytesIO

st.set_page_config(page_title="Report - K's Alpha Calculator", layout="wide")

st.title("Step 7: Generate Report")

# Check prerequisites
if not st.session_state.app_state.get('analysis_complete'):
    st.warning("Please run the analysis on the Analysis page first.")
    st.stop()

coder_data = st.session_state.app_state['coder_data']
per_var = st.session_state.app_state['per_variable_results']
pairwise_overall = st.session_state.app_state['pairwise_overall']
coder_impact = st.session_state.app_state['coder_impact']
overall_alpha = st.session_state.app_state['overall_alpha']
disagreements = st.session_state.app_state['disagreements']

st.markdown("""
Generate a comprehensive reliability report grounded in Krippendorff's methodology.
This report can be used in your methods section or as supplementary materials.
""")

st.markdown("---")

# Study metadata
st.subheader("Study Information")

col1, col2 = st.columns(2)

with col1:
    study_name = st.text_input("Study/Project Name", value="Intercoder Reliability Analysis")
    codebook_name = st.text_input("Codebook Name", value="Content Analysis Codebook")
    num_coders = st.number_input("Number of Coders", value=coder_data.n_coders, disabled=True)
    num_units = st.number_input("Number of Units Coded", value=coder_data.n_units, disabled=True)

with col2:
    coding_procedure = st.text_area(
        "Coding Procedure Description",
        value="Multiple coders independently coded each unit according to the codebook.",
        height=100,
    )
    training_description = st.text_area(
        "Coder Training Description",
        value="Coders were trained on the codebook and completed practice coding before the main study.",
        height=100,
    )

st.markdown("---")

# Report sections
st.subheader("Report Sections")

sections = st.multiselect(
    "Select sections to include:",
    [
        "Executive Summary",
        "Methodology Overview",
        "Overall Reliability",
        "Per-Variable Analysis",
        "Pairwise Coder Comparison",
        "Coder Impact Analysis",
        "Problem Areas & Recommendations",
        "Technical Details",
    ],
    default=[
        "Executive Summary",
        "Overall Reliability",
        "Per-Variable Analysis",
        "Problem Areas & Recommendations",
    ],
)

# Theory context toggle
use_theory = st.checkbox(
    "Include theory context from Krippendorff's book",
    value=True,
    help="Extracts relevant passages from 'Content Analysis: An Introduction to Its Methodology'",
)

st.markdown("---")

# Generate report
if st.button("Generate Report", type="primary", use_container_width=True):
    with st.spinner("Generating comprehensive report..."):
        try:
            from core.krippendorff import results_summary, interpret_alpha

            # Build results summary
            summary = results_summary(per_var, pairwise_overall, coder_impact)

            # Study metadata
            metadata = {
                "Study Name": study_name,
                "Codebook": codebook_name,
                "Number of Coders": str(coder_data.n_coders),
                "Number of Units": str(coder_data.n_units),
                "Number of Variables": str(len(coder_data.variables)),
                "Coding Procedure": coding_procedure,
                "Training": training_description,
            }

            # Extract theory context if requested
            theory_context = ""
            if use_theory:
                try:
                    import pdfplumber

                    book_path = "/Users/cemalatas/Desktop/vscode/kalpha_calculator/sources/Content Analysis An Introduction to Its Methodology (Klaus Krippendorff).pdf"

                    with pdfplumber.open(book_path) as pdf:
                        # Extract relevant pages (Chapter 11 on reliability, approximately pages 277-318)
                        relevant_pages = list(range(276, 290))  # Adjust based on actual content
                        theory_text = ""

                        for page_num in relevant_pages[:5]:  # Limit to avoid token overflow
                            if page_num < len(pdf.pages):
                                page = pdf.pages[page_num]
                                text = page.extract_text()
                                if text:
                                    theory_text += text[:2000] + "\n\n"  # Limit per page

                        theory_context = theory_text[:5000]  # Limit total

                except Exception as e:
                    st.warning(f"Could not load theory content: {e}")
                    theory_context = ""

            # Generate report using LLM
            if st.session_state.app_state.get('llm_client'):
                from core.llm_client import generate_report

                report = generate_report(
                    st.session_state.app_state['llm_client'],
                    summary,
                    metadata,
                    theory_context,
                )
            else:
                # Fallback: generate basic report without LLM
                report = generate_basic_report(
                    metadata, per_var, pairwise_overall, coder_impact, overall_alpha
                )

            st.session_state.app_state['generated_report'] = report
            st.success("Report generated successfully!")

        except Exception as e:
            st.error(f"Report generation failed: {str(e)}")
            import traceback
            st.code(traceback.format_exc())


def generate_basic_report(metadata, per_var, pairwise_overall, coder_impact, overall_alpha):
    """Generate a basic report without LLM."""
    from core.krippendorff import interpret_alpha

    lines = [
        f"# Intercoder Reliability Report",
        f"## {metadata.get('Study Name', 'Study')}",
        "",
        "---",
        "",
        "## Executive Summary",
        "",
        f"This report presents the intercoder reliability analysis for {metadata.get('Codebook', 'the study')}.",
        f"**{metadata.get('Number of Coders', 'N')} coders** coded **{metadata.get('Number of Units', 'N')} units** across **{metadata.get('Number of Variables', 'N')} variables**.",
        "",
        f"**Overall Krippendorff's Alpha: {overall_alpha:.3f}** ({interpret_alpha(overall_alpha)[0]})",
        "",
        "---",
        "",
        "## Per-Variable Reliability",
        "",
        "| Variable | Alpha | Status |",
        "|----------|-------|--------|",
    ]

    for r in per_var:
        interp, _ = interpret_alpha(r.alpha_all_coders)
        lines.append(f"| {r.variable} | {r.alpha_all_coders:.3f} | {interp} |")

    lines.extend([
        "",
        "---",
        "",
        "## Interpretation Guidelines",
        "",
        "- **α ≥ 0.80**: Acceptable reliability - proceed with analysis",
        "- **0.67 ≤ α < 0.80**: Tentative reliability - report with caution",
        "- **α < 0.67**: Insufficient reliability - consider revising codebook or retraining coders",
        "",
        "*Based on Krippendorff, K. (2019). Content Analysis: An Introduction to Its Methodology.*",
    ])

    return "\n".join(lines)


# Display report
if st.session_state.app_state.get('generated_report'):
    st.markdown("---")
    st.subheader("Generated Report")

    report = st.session_state.app_state['generated_report']

    # Display in expandable container
    with st.expander("View Report", expanded=True):
        st.markdown(report)

    # Export options
    st.markdown("---")
    st.subheader("Export Report")

    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        # Markdown
        st.download_button(
            "Download Markdown",
            report,
            file_name="reliability_report.md",
            mime="text/markdown",
        )

    with col2:
        # DOCX export
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            doc.add_heading('Intercoder Reliability Report', 0)

            # Split by sections and add
            for line in report.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)

            # Save to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            st.download_button(
                "Download DOCX",
                docx_buffer,
                file_name="reliability_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except ImportError:
            st.button("Download DOCX", disabled=True, help="Install python-docx for DOCX export")

    with col3:
        # CSV - Per variable results
        per_var_df = pd.DataFrame([r.to_dict() for r in per_var])
        csv_buffer = per_var_df.to_csv(index=False)
        st.download_button(
            "Results CSV",
            csv_buffer,
            file_name="reliability_results.csv",
            mime="text/csv",
        )

    with col4:
        # XLSX - All results
        try:
            xlsx_buffer = BytesIO()
            with pd.ExcelWriter(xlsx_buffer, engine='openpyxl') as writer:
                # Per-variable sheet
                per_var_df = pd.DataFrame([r.to_dict() for r in per_var])
                per_var_df.to_excel(writer, sheet_name='Per-Variable', index=False)

                # Pairwise sheet
                pairwise_df = pd.DataFrame([p.to_dict() for p in pairwise_overall])
                pairwise_df.to_excel(writer, sheet_name='Pairwise', index=False)

                # Coder impact sheet
                impact_df = pd.DataFrame([c.to_dict() for c in coder_impact])
                impact_df.to_excel(writer, sheet_name='Coder Impact', index=False)

                # Disagreements sheet
                if disagreements:
                    disag_df = pd.DataFrame([d.to_dict() for d in disagreements[:1000]])
                    disag_df.to_excel(writer, sheet_name='Disagreements', index=False)

            xlsx_buffer.seek(0)

            st.download_button(
                "All Results XLSX",
                xlsx_buffer,
                file_name="reliability_results.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except Exception as e:
            st.button("All Results XLSX", disabled=True, help=f"Error: {e}")

    with col5:
        # PDF export
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet

            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
            styles = getSampleStyleSheet()

            story = []
            for line in report.split('\n'):
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))

            doc.build(story)
            pdf_buffer.seek(0)

            st.download_button(
                "Download PDF",
                pdf_buffer,
                file_name="reliability_report.pdf",
                mime="application/pdf",
            )
        except ImportError:
            st.button("Download PDF", disabled=True, help="Install reportlab for PDF export")

st.markdown("---")
st.info("Report generation complete! You can download your results in multiple formats above.")
